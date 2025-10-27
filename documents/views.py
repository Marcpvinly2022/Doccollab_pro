from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt
import json
from .models import Document, DocumentPermission, UserPresence, DocumentVersion, DocumentComment, DocumentActivity
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
import html2text
import io
import re
from html import unescape
def register(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')
        password_confirm = request.POST.get('password_confirm')
        
        if password != password_confirm:
            return render(request, 'register.html', {'error': 'Passwords do not match'})
        
        if User.objects.filter(username=username).exists():
            return render(request, 'register.html', {'error': 'Username already exists'})
        
        user = User.objects.create_user(username=username, email=email, password=password)
        login(request, user)
        return redirect('dashboard')
    
    return render(request, 'register.html')

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('dashboard')
        else:
            return render(request, 'login.html', {'error': 'Invalid credentials'})
    
    return render(request, 'login.html')

def logout_view(request):
    logout(request)
    return redirect('login')

@login_required(login_url='login')
def dashboard(request):
    owned_documents = Document.objects.filter(owner=request.user)
    shared_documents = Document.objects.filter(permissions__user=request.user).distinct()
    
    context = {
        'owned_documents': owned_documents,
        'shared_documents': shared_documents,
    }
    return render(request, 'dashboard.html', context)

@login_required(login_url='login')
def create_document(request):
    if request.method == 'POST':
        title = request.POST.get('title', 'Untitled Document')
        document = Document.objects.create(
            title=title,
            owner=request.user,
            content={"type": "doc", "content": []}
        )
        return redirect('editor', document_id=document.id)
    
    return render(request, 'create_document.html')

@login_required(login_url='login')
def editor(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    # Check permissions
    if document.owner != request.user and not document.is_public:
        if not document.permissions.filter(user=request.user).exists():
            return redirect('dashboard')
    
    active_users = UserPresence.objects.filter(document=document).select_related('user')
    comments = DocumentComment.objects.filter(document=document)
    activities = DocumentActivity.objects.filter(document=document)[:20]
    
    context = {
        'document': document,
        'active_users': active_users,
        'comments': comments,
        'activities': activities,
        'is_owner': document.owner == request.user,
        'can_edit': document.owner == request.user or document.permissions.filter(user=request.user, permission='editor').exists(),
    }
    return render(request, 'editor.html', context)


"""
@login_required(login_url='login')
@require_http_methods(["POST"])
def share_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    if document.owner != request.user:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    data = json.loads(request.body)
    username = data.get('username')
    permission = data.get('permission', 'editor')
    
    try:
        user = User.objects.get(username=username)
        DocumentPermission.objects.update_or_create(
            document=document,
            user=user,
            defaults={'permission': permission}
        )
        DocumentActivity.objects.create(
            document=document,
            user=request.user,
            activity_type='share',
            description=f'Shared with {username} ({permission})'
        )
        return JsonResponse({'success': True})
    except User.DoesNotExist:
        return JsonResponse({'error': 'User not found'}, status=404)
"""

@login_required(login_url='login')
@require_http_methods(["POST"])
def share_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    if document.owner != request.user:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    data = json.loads(request.body)
    username = data.get('username')
    permission = data.get('permission', 'editor')
    
    # DEBUG: Print all users for testing
    all_users = User.objects.all().values('username', 'email')
    print("All registered users:", list(all_users))
    
    try:
        # FIX 1: Case-insensitive lookup or exact match
        user = User.objects.get(username__iexact=username)  # Case-insensitive
        
        # FIX 2: Check if trying to share with yourself
        if user == request.user:
            return JsonResponse({'error': 'Cannot share document with yourself'}, status=400)
        
        # FIX 3: Check if user already has permission
        existing_permission = DocumentPermission.objects.filter(
            document=document, 
            user=user
        ).first()
        
        if existing_permission:
            return JsonResponse({
                'error': f'Document already shared with {username}'
            }, status=400)
        
        # Create permission
        DocumentPermission.objects.create(
            document=document,
            user=user,
            permission=permission
        )
        
        # Log activity
        DocumentActivity.objects.create(
            document=document,
            user=request.user,
            activity_type='share',
            description=f'Shared with {username} ({permission})'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Document successfully shared with {username}'
        })
        
    except User.DoesNotExist:
        # FIX 4: Better error reporting
        return JsonResponse({
            'error': f'User "{username}" not found. Available users: {list(User.objects.values_list("username", flat=True))}'
        }, status=404)




@login_required(login_url='login')
@require_http_methods(["POST"])
def toggle_public(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    if document.owner != request.user:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    document.is_public = not document.is_public
    document.save()
    return JsonResponse({'is_public': document.is_public})

@login_required(login_url='login')
@require_http_methods(["DELETE"])
def delete_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    if document.owner != request.user:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    document.delete()
    return JsonResponse({'success': True})

@login_required(login_url='login')
@require_http_methods(["POST"])
def save_version(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    if document.owner != request.user and not document.permissions.filter(user=request.user, permission='editor').exists():
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    data = json.loads(request.body)
    version_count = DocumentVersion.objects.filter(document=document).count()
    
    DocumentVersion.objects.create(
        document=document,
        content=document.content,
        created_by=request.user,
        version_number=version_count + 1,
        change_summary=data.get('summary', 'Manual save')
    )
    
    return JsonResponse({'success': True, 'version': version_count + 1})

@login_required(login_url='login')
def get_versions(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    if document.owner != request.user and not document.is_public:
        if not document.permissions.filter(user=request.user).exists():
            return JsonResponse({'error': 'Permission denied'}, status=403)
    
    versions = DocumentVersion.objects.filter(document=document).values(
        'id', 'version_number', 'created_by__username', 'created_at', 'change_summary'
    )
    
    return JsonResponse({'versions': list(versions)})

@login_required(login_url='login')
def restore_version(request, document_id, version_id):
    document = get_object_or_404(Document, id=document_id)
    version = get_object_or_404(DocumentVersion, id=version_id, document=document)
    
    if document.owner != request.user and not document.permissions.filter(user=request.user, permission='editor').exists():
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    document.content = version.content
    document.save()
    
    DocumentActivity.objects.create(
        document=document,
        user=request.user,
        activity_type='edit',
        description=f'Restored to version {version.version_number}'
    )
    
    return JsonResponse({'success': True})

"""
@csrf_exempt
def download_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    format_type = request.GET.get('format', 'txt')
    
    # Use content from POST request if available, otherwise use saved content
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            content = data.get('content', document.content)
            title = data.get('title', document.title)
        except json.JSONDecodeError:
            content = document.content
            title = document.title
    else:
        content = document.content
        title = document.title
    
    if format_type == 'txt':
        return generate_txt(content, title)
    elif format_type == 'pdf':
        return generate_pdf(content, title)
    elif format_type == 'docx':
        return generate_docx(content, title)
    else:
        return JsonResponse({'error': 'Unsupported format'}, status=400)

def generate_txt(content, title):
    # Convert HTML to plain text
    import re
    from html import unescape
    
    # Remove HTML tags and convert entities
    text = re.sub('<[^<]+?>', '', content)
    text = unescape(text)
    
    response = HttpResponse(f"{title}\n\n{text}", content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="{title}.txt"'
    return response

def generate_pdf(content, title):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer)
    
    # Set up PDF content
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, 800, title)
    
    # Convert HTML to plain text for PDF
    import re
    from html import unescape
    text = re.sub('<[^<]+?>', '', content)
    text = unescape(text)
    
    # Simple content rendering
    p.setFont("Helvetica", 12)
    y_position = 750
    lines = text.split('\n')
    
    for line in lines:
        if y_position < 50:  # Add new page if needed
            p.showPage()
            p.setFont("Helvetica", 12)
            y_position = 800
        
        # Handle long lines by wrapping
        if len(line) > 80:
            words = line.split(' ')
            current_line = ""
            for word in words:
                if len(current_line + word) < 80:
                    current_line += word + " "
                else:
                    p.drawString(50, y_position, current_line.strip())
                    y_position -= 20
                    current_line = word + " "
                    if y_position < 50:
                        p.showPage()
                        p.setFont("Helvetica", 12)
                        y_position = 800
            if current_line:
                p.drawString(50, y_position, current_line.strip())
                y_position -= 20
        else:
            p.drawString(50, y_position, line)
            y_position -= 20
        
        if y_position < 50:
            p.showPage()
            p.setFont("Helvetica", 12)
            y_position = 800
    
    p.showPage()
    p.save()
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{title}.pdf"'
    return response

def generate_docx(content, title):
    doc = Document()
    
    # Add title
    doc.add_heading(title, 0)
    
    # Convert HTML to plain text for DOCX
    import re
    from html import unescape
    text = re.sub('<[^<]+?>', '', content)
    text = unescape(text)
    
    # Add content to Word document
    paragraphs = text.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():  # Only add non-empty paragraphs
            doc.add_paragraph(paragraph)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer, 
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
    return response






@login_required(login_url='login')
def download_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    
    if document.owner != request.user and not document.is_public:
        if not document.permissions.filter(user=request.user).exists():
            return JsonResponse({'error': 'Permission denied'}, status=403)
    
    file_format = request.GET.get('format', 'txt')
    
    if file_format == 'docx':
        return download_as_docx(document)
    elif file_format == 'pdf':
        return download_as_pdf(document)
    else:
        return download_as_txt(document)

def download_as_txt(document):
    content = extract_text_from_tiptap(document.content)
    response = HttpResponse(content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="{document.title}.txt"'
    return response

def download_as_docx(document):
    doc = DocxDocument()
    doc.add_heading(document.title, 0)
    content = extract_text_from_tiptap(document.content)
    doc.add_paragraph(content)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
    return response

def download_as_pdf(document):
    content = extract_text_from_tiptap(document.content)
    
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 750, document.title)
    
    c.setFont("Helvetica", 12)
    y = 720
    for line in content.split('\n'):
        if y < 50:
            c.showPage()
            y = 750
        c.drawString(50, y, line[:80])
        y -= 20
    
    c.save()
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{document.title}.pdf"'
    return response

def extract_text_from_tiptap(content):
    #Extract plain text from Tiptap JSON content
    if isinstance(content, dict):
        text = []
        if 'content' in content:
            for node in content['content']:
                text.append(extract_text_from_node(node))
        return '\n'.join(text)
    return ''

def extract_text_from_node(node):
    #Recursively extract text from a Tiptap node
    if node.get('type') == 'text':
        return node.get('text', '')
    
    text = []
    if 'content' in node:
        for child in node['content']:
            text.append(extract_text_from_node(child))
    
    return ''.join(text)

"""




@csrf_exempt
def download_document(request, document_id):
    from .models import Document
    
    try:
        document = get_object_or_404(Document, id=document_id)
        format_type = request.GET.get('format', 'txt')
        
        # Get content from request or use saved content
        if request.method == 'POST':
            try:
                data = json.loads(request.body)
                content = data.get('content', '')
                title = data.get('title', document.title)
            except:
                content = document.content
                title = document.title
        else:
            content = document.content
            title = document.title
        
        if format_type == 'txt':
            return generate_txt(content, title)
        elif format_type == 'pdf':
            return generate_pdf(content, title)
        elif format_type == 'docx':
            return generate_docx(content, title)
        else:
            return JsonResponse({'error': 'Unsupported format'}, status=400)
            
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def generate_txt(content, title):
    """Generate plain text file"""
    # Simple HTML to text conversion
    import re
    from html import unescape
    
    # Remove HTML tags
    text = re.sub('<[^<]+?>', '', content)
    text = unescape(text)
    
    response = HttpResponse(f"{title}\n\n{text}", content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="{title}.txt"'
    return response

def generate_pdf(content, title):
    """Generate PDF file - Simple version"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        
        # Add title
        p.setFont("Helvetica-Bold", 16)
        p.drawString(100, 750, title)
        
        # Convert HTML to text
        import re
        from html import unescape
        text = re.sub('<[^<]+?>', '', content)
        text = unescape(text)
        
        # Add content
        p.setFont("Helvetica", 12)
        y = 700
        lines = text.split('\n')
        
        for line in lines:
            if y < 100:  # New page if needed
                p.showPage()
                p.setFont("Helvetica", 12)
                y = 750
            p.drawString(50, y, line[:80])  # Limit line length
            y -= 20
        
        p.save()
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{title}.pdf"'
        return response
        
    except ImportError:
        return JsonResponse({'error': 'PDF generation requires reportlab package'}, status=500)
    except Exception as e:
        return JsonResponse({'error': f'PDF generation failed: {str(e)}'}, status=500)

def generate_docx(content, title):
    """Generate DOCX file - Simple version"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # Convert HTML to text
        import re
        from html import unescape
        text = re.sub('<[^<]+?>', '', content)
        text = unescape(text)
        
        # Add content
        lines = text.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
        return response
        
    except ImportError:
        return JsonResponse({'error': 'DOCX generation requires python-docx package'}, status=500)
    except Exception as e:
        return JsonResponse({'error': f'DOCX generation failed: {str(e)}'}, status=500)

